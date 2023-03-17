import docx

def create_word(fio, comp, doc, car_num):

    doc = docx.Document("C:/Users/povar/OneDrive/Desktop/pass_bot/wq.docx")


    for p in doc.paragraphs:
        # Проверяем, содержит ли абзац заданное слово
        if 'Ф., И., О.,' in p.text:
            # Ищем позицию заданного слова в тексте абзаца
            pos = p.text.find('Ф., И., О.,')
            # Создаем новый объект "Run" с необходимым текстом
            new_run = p.add_run(fio)
            # Изменяем оформление нового объекта "Run", если нужно
            new_run.bold = True
            # Разбиваем текст исходного абзаца на две части и вставляем новый объект "Run" между ними
            p.text = p.text[:pos + len(fio)] + p.text[pos + len(fio):]

    for p in doc.paragraphs:    
        if 'Наименование организации' in p.text:
            # Ищем позицию заданного слова в тексте абзаца
            pos = p.text.find('Наименование организации')
            # Создаем новый объект "Run" с необходимым текстом
            new_run = p.add_run(doc)
            # Изменяем оформление нового объекта "Run", если нужно
            new_run.bold = True
            # Разбиваем текст исходного абзаца на две части и вставляем новый объект "Run" между ними
            p.text = p.text[:pos + len(doc)] + p.text[pos + len(doc):]
    
    # for p in doc.paragraphs:    
    #     if 'Документ водителя' in p.text:
    #         # Ищем позицию заданного слова в тексте абзаца
    #         pos = p.text.find('Документ водителя')
    #         # Создаем новый объект "Run" с необходимым текстом
    #         new_run = p.add_run(comp)
    #         # Изменяем оформление нового объекта "Run", если нужно
    #         new_run.bold = True
    #         # Разбиваем текст исходного абзаца на две части и вставляем новый объект "Run" между ними
    #         p.text = p.text[:pos + len(comp)] + p.text[pos + len(comp):]


    for p in doc.paragraphs:
        if 'Номер а/м' in p.text:
            # Ищем позицию заданного слова в тексте абзаца
            pos = p.text.find('Номер а/м')
            # Создаем новый объект "Run" с необходимым текстом
            new_run = p.add_run(car_num)
            # Изменяем оформление нового объекта "Run", если нужно
            new_run.bold = True
            # Разбиваем текст исходного абзаца на две части и вставляем новый объект "Run" между ними
            p.text = p.text[:pos + len(car_num)] + p.text[pos + len(car_num):]


    # Сохраняем документ
    doc.save('пропуск.docx')
